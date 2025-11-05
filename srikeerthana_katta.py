import streamlit as st
import pytesseract
from PIL import Image
import cv2
import numpy as np
import random
import uuid
from datetime import datetime
import requests
import os
import tempfile

# -------------------- PDF & DOC PROCESSING IMPORTS --------------------
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    try:
        import pypdf
        PDF_AVAILABLE = True
        PyPDF2 = pypdf
    except ImportError:
        PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Set Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# -------------------- PAGE CONFIG --------------------
st.set_page_config(
    page_title="ChatGPT",
    layout="wide",
    initial_sidebar_state="expanded"
)

# -------------------- CUSTOM CSS (Modern UI) --------------------
st.markdown("""
<style>
/* Main background */
[data-testid="stAppViewContainer"] {
    background-color: #343541;
}
/* Sidebar */
[data-testid="stSidebar"] {
    background-color: #202123;
}
.stButton>button {
    background-color: transparent;
    color: white;
    border: 1px solid #565869;
    border-radius: 6px;
    margin-bottom: 6px;
}
.stButton>button:hover {
    background-color: #2b2c2f;
}
/* Chat bubbles */
.chat-message {
    padding: 12px;
    border-radius: 12px;
    margin: 6px 0;
    max-width: 80%;
}
.user-message {
    background-color: #10a37f;
    color: white;
    margin-left: auto;
}
.assistant-message {
    background-color: #40414f;
    color: white;
    margin-right: auto;
}
/* Welcome cards */
.card {
    background: rgba(255,255,255,0.05);
    border-radius: 12px;
    padding: 16px;
}
.card b {
    color: white;
}
.card div {
    color:#8e8ea0;
    font-size:14px;
}
/* Session ID container */
.session-id-container {
    background-color: rgba(255,255,255,0.05);
    border-radius: 8px;
    padding: 12px;
    margin-bottom: 16px;
    border: 1px solid rgba(255,255,255,0.1);
}
.session-id-label {
    font-size: 12px;
    color: #8e8ea0;
    margin-bottom: 4px;
}
.session-id-value {
    font-family: 'Courier New', monospace;
    font-size: 14px;
    color: #10a37f;
    font-weight: bold;
}
.session-time {
    font-size: 11px;
    color: #8e8ea0;
    margin-top: 4px;
}
/* File upload styling */
.uploaded-file {
    background: rgba(255,255,255,0.05);
    border-radius: 8px;
    padding: 12px;
    margin: 8px 0;
    border: 1px solid rgba(255,255,255,0.1);
}
.file-icon {
    font-size: 24px;
    margin-right: 10px;
}
.file-info {
    display: flex;
    align-items: center;
    margin-bottom: 8px;
}
.file-content {
    background: rgba(0,0,0,0.3);
    border-radius: 6px;
    padding: 10px;
    max-height: 200px;
    overflow-y: auto;
    font-size: 12px;
    line-height: 1.4;
}
.file-type-badge {
    display: inline-block;
    background: rgba(16, 163, 127, 0.2);
    color: #10a37f;
    padding: 4px 8px;
    border-radius: 4px;
    font-size: 12px;
    margin: 2px;
}
/* Hide Streamlit branding */
#MainMenu, header, footer {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# -------------------- SESSION --------------------
if "session_id" not in st.session_state:
    st.session_state["session_id"] = str(uuid.uuid4())[:8]
if "session_start_time" not in st.session_state:
    st.session_state["session_start_time"] = datetime.now()
if "messages" not in st.session_state:
    st.session_state["messages"] = []
if "ollama_enabled" not in st.session_state:
    st.session_state["ollama_enabled"] = False
if "uploaded_files" not in st.session_state:
    st.session_state["uploaded_files"] = []

# -------------------- DOCUMENT PROCESSING FUNCTIONS --------------------
def extract_text_from_image(image):
    """Extract text from image using OCR"""
    try:
        img_array = np.array(image)
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        gray = cv2.medianBlur(gray, 3)
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        text = pytesseract.image_to_string(thresh, config='--oem 3 --psm 6')
        return text.strip() if text.strip() else "No text found in the image."
    except Exception as e:
        return f"OCR Error: {str(e)}"

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    if not PDF_AVAILABLE:
        return "PDF processing unavailable. Please install PyPDF2: pip install pypdf"
    
    try:
        # Save uploaded file to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(pdf_file.getvalue())
            tmp_path = tmp_file.name
        
        # Read from temporary file
        with open(tmp_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # Clean up
        os.unlink(tmp_path)
        
        return text.strip() if text.strip() else "No extractable text found in PDF."
    except Exception as e:
        return f"PDF Processing Error: {str(e)}"

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    if not DOCX_AVAILABLE:
        return "DOCX processing unavailable. Install: pip install python-docx"
    
    try:
        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(docx_file.getvalue())
            tmp_path = tmp_file.name
        
        # Read from temporary file
        doc = Document(tmp_path)
        full_text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        text = "\n".join(full_text)
        
        # Clean up
        os.unlink(tmp_path)
        
        return text.strip() if text.strip() else "No text found in DOCX file."
    except Exception as e:
        return f"DOCX Processing Error: {str(e)}"

def extract_text_from_txt(txt_file):
    """Extract text from TXT file"""
    try:
        txt_file.seek(0)
        text = txt_file.read().decode('utf-8')
        return text.strip() if text.strip() else "No text found in TXT file."
    except Exception as e:
        return f"TXT Processing Error: {str(e)}"

def process_uploaded_file(uploaded_file):
    """Process any uploaded file and extract text"""
    file_name = uploaded_file.name.lower()
    
    if file_name.endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(uploaded_file)
        text = extract_text_from_image(image)
        file_icon = "🖼"
        file_type_name = "Image"
        
    elif file_name.endswith('.pdf'):
        text = extract_text_from_pdf(uploaded_file)
        file_icon = "📄"
        file_type_name = "PDF"
        
    elif file_name.endswith(('.docx', '.doc')):
        text = extract_text_from_docx(uploaded_file)
        file_icon = "📝"
        file_type_name = "Word Document"
        
    elif file_name.endswith('.txt'):
        text = extract_text_from_txt(uploaded_file)
        file_icon = "📃"
        file_type_name = "Text File"
        
    else:
        text = f"Unsupported file type: {uploaded_file.type}"
        file_icon = "❓"
        file_type_name = "Unknown"
    
    return {
        "name": uploaded_file.name,
        "type": file_type_name,
        "icon": file_icon,
        "content": text,
        "timestamp": datetime.now().strftime("%H:%M:%S"),
        "size": f"{uploaded_file.size / 1024:.1f} KB"
    }

# -------------------- OLLAMA FUNCTION --------------------
def ollama_available():
    try:
        r = requests.get("http://localhost:11434/api/tags", timeout=1)
        return r.status_code == 200
    except:
        return False

def ollama_response(prompt):
    try:
        data = {
            "model": "llama3.2",
            "prompt": prompt,
            "stream": False
        }
        r = requests.post("http://localhost:11434/api/generate", json=data, timeout=30)
        if r.status_code == 200:
            return r.json().get("response", "No response from Ollama.")
    except Exception as e:
        return f"Ollama Error: {str(e)}"
    
    return random.choice([
        "Hello! How can I assist today?",
        "I'm here to help you with whatever you need.",
        "Got it — let's work on that."
    ])

# -------------------- SIDEBAR --------------------
with st.sidebar:
    st.markdown(f"""
    <div class="session-id-container">
        <div class="session-id-label">SESSION ID</div>
        <div class="session-id-value">{st.session_state["session_id"]}</div>
        <div class="session-time">Started: {st.session_state["session_start_time"].strftime("%H:%M:%S")}</div>
    </div>
    """, unsafe_allow_html=True)

    st.button("➕  New chat")
    st.text_input("Search chats", label_visibility="collapsed", placeholder="Search chats...")
    st.button("📚 Library")
    st.button("🛠 Projects")

    st.markdown("#### Chats")
    for chat in ["Packing checklist", "Bible quiz prompt", "Infosys Internship"]:
        st.button(chat, use_container_width=True)

    st.markdown("---")
    
    # Enhanced File Upload Section
    st.markdown("#### 📁 Upload Documents")
    
    # File type information
    st.markdown("""
    <div style='background: rgba(255,255,255,0.05); padding: 12px; border-radius: 8px; margin-bottom: 12px;'>
        <div style="font-size: 14px; color: white; margin-bottom: 8px;">Supported File Types:</div>
        <span class="file-type-badge">PDF</span>
        <span class="file-type-badge">DOC/DOCX</span>
        <span class="file-type-badge">TXT</span>
        <span class="file-type-badge">PNG/JPG</span>
    </div>
    """, unsafe_allow_html=True)
    
    # File uploader with ALL document types
    uploaded_files = st.file_uploader(
        "Choose files",
        type=['pdf', 'docx', 'doc', 'txt', 'png', 'jpg', 'jpeg'],
        accept_multiple_files=True,
        label_visibility="collapsed",
        help="Select PDFs, Word documents, text files, or images"
    )
    
    # Process uploaded files
    if uploaded_files:
        for uploaded_file in uploaded_files:
            # Check if file is already processed
            existing_files = [f["name"] for f in st.session_state["uploaded_files"]]
            if uploaded_file.name not in existing_files:
                with st.spinner(f"Processing {uploaded_file.name}..."):
                    file_data = process_uploaded_file(uploaded_file)
                    st.session_state["uploaded_files"].append(file_data)
                st.success(f"✅ {uploaded_file.name}")
    
    # Display uploaded files in sidebar
    if st.session_state["uploaded_files"]:
        st.markdown("#### Uploaded Files")
        for file_data in st.session_state["uploaded_files"][-5:]:  # Show last 5 files
            with st.expander(f"{file_data['icon']} {file_data['name']}"):
                st.markdown(f"*Type:* {file_data['type']}")
                st.markdown(f"*Size:* {file_data['size']}")
                st.markdown(f"*Uploaded:* {file_data['timestamp']}")
                st.markdown("*Content Preview:*")
                preview_text = file_data['content'][:500] + "..." if len(file_data['content']) > 500 else file_data['content']
                st.text_area("", value=preview_text, height=150, key=f"preview_{file_data['name']}", label_visibility="collapsed")

    st.markdown("---")
    if st.toggle("🤖 Ollama Model"):
        st.session_state["ollama_enabled"] = True
        if not ollama_available():
            st.warning("Ollama is not running on localhost:11434")
    else:
        st.session_state["ollama_enabled"] = False

# -------------------- CHAT UI --------------------
st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
st.title("Good to see you, srikkeerthana 👋")

# Display messages
for message in st.session_state["messages"]:
    role_class = "user-message" if message["role"] == "user" else "assistant-message"
    st.markdown(f"""
    <div class="chat-message {role_class}">
        {message["content"]}
    </div>
    """, unsafe_allow_html=True)

# -------------------- UPLOADED DOCUMENTS DISPLAY --------------------
if st.session_state["uploaded_files"]:
    st.markdown("---")
    st.subheader("📂 Your Uploaded Documents")
    
    for file_data in st.session_state["uploaded_files"]:
        st.markdown(f"""
        <div class="uploaded-file">
            <div class="file-info">
                <span class="file-icon">{file_data['icon']}</span>
                <div>
                    <strong>{file_data['name']}</strong><br>
                    <small>Type: {file_data['type']} • Size: {file_data['size']} • Uploaded: {file_data['timestamp']}</small>
                </div>
            </div>
            <div class="file-content">
                {file_data['content'][:300]}{'...' if len(file_data['content']) > 300 else ''}
            </div>
        </div>
        """, unsafe_allow_html=True)

# -------------------- Chat Input --------------------
prompt = st.chat_input("Ask anything or discuss uploaded documents...")

if prompt:
    st.session_state["messages"].append({"role": "user", "content": prompt})
    
    # Include uploaded file content in context
    context = ""
    if st.session_state["uploaded_files"]:
        context = "\n\nReference from uploaded documents:\n"
        for file_data in st.session_state["uploaded_files"]:
            context += f"\n--- {file_data['name']} ({file_data['type']}) ---\n"
            context += file_data['content'][:800] + ("..." if len(file_data['content']) > 800 else "")
    
    full_prompt = prompt + context
    
    if st.session_state["ollama_enabled"] and ollama_available():
        reply = ollama_response(full_prompt)
    else:
        if st.session_state["uploaded_files"]:
            doc_responses = [
                f"I've analyzed your {len(st.session_state['uploaded_files'])} uploaded document(s). Regarding '{prompt}', I can help you work with the content.",
                f"Based on your documents, I can assist you with '{prompt}'. The files contain relevant information we can discuss.",
                f"I see you've uploaded {len(st.session_state['uploaded_files'])} document(s). Let me help you analyze them in relation to: '{prompt}'"
            ]
            reply = random.choice(doc_responses)
        else:
            reply = random.choice([
                "Hello! How can I assist today?",
                "I'm here to help you with whatever you need!",
                "Let's get started 🚀",
                "What would you like to explore today?",
                "Ready to help! What's on your mind?"
            ])
    
    st.session_state["messages"].append({"role": "assistant", "content": reply})
    st.rerun()

# -------------------- Welcome UI --------------------
if not st.session_state["messages"]:
    st.markdown("""
    <div style="text-align: center; margin-top: 100px;">
      <h1 style="background: linear-gradient(90deg, #5436DA, #16B3EF); 
      -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-size: 32px;">
      ChatGPT with Document Support</h1>
      <p style="color: #d1d5db;">Upload PDFs, Word docs, images, or text files and chat about their content!</p>
      <div style="display:grid; grid-template-columns:repeat(3, 1fr); gap:16px; margin-top: 40px;">
        <div class="card"><b>📄 PDF Analysis</b><div>Upload and discuss PDF documents</div></div>
        <div class="card"><b>📝 Word Docs</b><div>Process DOC/DOCX files</div></div>
        <div class="card"><b>🖼 Image OCR</b><div>Text extraction from images</div></div>
      </div>
      <div style="display:grid; grid-template-columns:repeat(3, 1fr); gap:16px; margin-top: 16px;">
        <div class="card"><b>📃 Text Files</b><div>Process .txt documents</div></div>
        <div class="card"><b>🤖 Local AI</b><div>Enable Ollama integration</div></div>
        <div class="card"><b>📁 Multi-file</b><div>Upload multiple files at once</div></div>
      </div>
    </div>
    """, unsafe_allow_html=True)

# -------------------- INSTALLATION INSTRUCTIONS --------------------
if not PDF_AVAILABLE or not DOCX_AVAILABLE:
    with st.expander("🔧 Installation Requirements"):
        st.markdown("### Required Packages")
        st.code("""
# Install all required packages:
pip install PyPDF2 python-docx pytesseract opencv-python pillow

# For OCR, also install Tesseract:
# Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki
# macOS: brew install tesseract
# Linux: sudo apt-get install tesseract-ocr
        """)
        
        if not PDF_AVAILABLE:
            st.error("PyPDF2 is not installed - PDF processing disabled")
        if not DOCX_AVAILABLE:
            st.error("python-docx is not installed - Word document processing disabled")